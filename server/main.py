from __future__ import annotations

import asyncio
import json
import os
import shutil
import uuid
from collections import defaultdict
from contextlib import asynccontextmanager
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Optional

import aiofiles
import httpx
import uvicorn
from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlalchemy import (
    Column,
    DateTime,
    Float,
    String,
    Text,
    create_engine,
    text,
)
from sqlalchemy.orm import Session, declarative_base, sessionmaker

from .pipeline import Pipeline, TaskProgress
from .templates import (
    MeetingType,
    TemplateCreateRequest,
    TemplateManager,
    TemplateRecord,
    TemplateUpdateRequest,
)


MAX_FILE_SIZE_BYTES = 500 * 1024 * 1024
ALLOWED_EXT = {".mp3", ".wav", ".m4a"}
WECHAT_WEBHOOK_URL = os.environ.get("WECHAT_WEBHOOK_URL", "")
VIEW_BASE_URL = os.environ.get("VIEW_BASE_URL", "http://localhost:8000")


Base = declarative_base()

DATABASE_URL = os.environ.get("DATABASE_URL", "sqlite:///./meetings.db")
UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "./uploads")
CHUNK_DIR = os.environ.get("CHUNK_DIR", "./chunks")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(CHUNK_DIR, exist_ok=True)

engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine, autocommit=False, autoflush=False)


class MeetingORM(Base):
    __tablename__ = "meetings"

    id = Column(String, primary_key=True)
    title = Column(String(500), default="")
    meeting_type = Column(String(50), default="client_consultation")
    case_id = Column(String(200), nullable=True)
    case_match_status = Column(String(50), default="pending")
    candidate_case_numbers = Column(Text, default="[]")
    participants = Column(Text, default="[]")
    case_background = Column(Text, default="")
    audio_filename = Column(String(500), default="")
    total_size_bytes = Column(Float, default=0.0)
    duration_seconds = Column(Float, default=0.0)
    transcript = Column(Text, default="")
    segments = Column(Text, default="[]")
    summary = Column(Text, default="{}")
    state = Column(String(50), default="uploading")
    error_message = Column(Text, default="")
    created_at = Column(DateTime, default=datetime.now)
    updated_at = Column(DateTime, default=datetime.now)


Base.metadata.create_all(bind=engine)
try:
    with engine.begin() as conn:
        conn.execute(text("ALTER TABLE meetings ADD COLUMN case_background TEXT DEFAULT ''"))
except Exception:
    pass
try:
    with engine.begin() as conn:
        conn.execute(text("ALTER TABLE meetings ADD COLUMN candidate_case_numbers TEXT DEFAULT '[]'"))
except Exception:
    pass
try:
    with engine.begin() as conn:
        conn.execute(text("ALTER TABLE meetings ADD COLUMN total_size_bytes REAL DEFAULT 0.0"))
except Exception:
    pass

template_manager = TemplateManager()

task_progress_store: dict[str, TaskProgress] = {}
task_ws_connections: dict[str, list[WebSocket]] = {}
task_queue: asyncio.Queue[str] = asyncio.Queue()

pipeline = Pipeline(
    template_manager=template_manager,
    whisper_model_size=os.environ.get("WHISPER_MODEL_SIZE", "medium"),
    openai_model=os.environ.get("OPENAI_MODEL", "gpt-4o"),
    openai_base_url=os.environ.get("OPENAI_BASE_URL"),
    openai_api_key=os.environ.get("OPENAI_API_KEY"),
    skip_diarization=os.environ.get("SKIP_DIARIZATION", "false").lower() == "true",
)


class MeetingCreateRequest(BaseModel):
    title: str = Field("", max_length=500)
    meeting_type: MeetingType = MeetingType.CLIENT_CONSULTATION
    participants: list[str] = Field(default_factory=list)
    case_background: Optional[str] = None


class MeetingResponse(BaseModel):
    id: str
    title: str
    meeting_type: str
    case_id: Optional[str] = None
    case_match_status: str = "pending"
    candidate_case_numbers: list[str] = []
    participants: list[str] = []
    case_background: str = ""
    audio_filename: str = ""
    total_size_bytes: float = 0.0
    duration_seconds: float = 0.0
    state: str = "uploading"
    error_message: str = ""
    created_at: str = ""
    updated_at: str = ""


class MeetingSummaryResponse(BaseModel):
    id: str
    title: str
    meeting_type: str
    participants: list[str] = []
    case_background: str = ""
    transcript: str = ""
    segments: list[dict[str, Any]] = []
    summary: dict[str, Any] = {}
    case_id: Optional[str] = None
    case_match_status: str = "pending"
    candidate_case_numbers: list[str] = []
    created_at: str = ""


class SearchResponse(BaseModel):
    results: list[MeetingSummaryResponse]
    total: int


class CaseMeetingItem(BaseModel):
    id: str
    title: str
    meeting_type: str
    participants: list[str] = []
    case_background: str = ""
    duration_seconds: float = 0.0
    todo_count: int = 0
    risk_count: int = 0
    candidate_case_numbers: list[str] = []
    case_id: Optional[str] = None
    case_match_status: str = "pending"
    state: str = ""
    created_at: str = ""


class CaseTimelineResponse(BaseModel):
    case_id: str
    meetings: list[CaseMeetingItem]
    total: int


class CaseSummaryItem(BaseModel):
    case_id: str
    case_match_status: str
    meeting_count: int = 0
    completed_count: int = 0
    todo_count: int = 0
    risk_count: int = 0
    latest_meeting_time: str = ""
    latest_meeting_title: str = ""


class CaseListResponse(BaseModel):
    cases: list[CaseSummaryItem]
    total: int


def _orm_to_response(m: MeetingORM) -> MeetingResponse:
    try:
        candidates = json.loads(m.candidate_case_numbers or "[]")
    except Exception:
        candidates = []
    return MeetingResponse(
        id=m.id,
        title=m.title,
        meeting_type=m.meeting_type,
        case_id=m.case_id,
        case_match_status=m.case_match_status,
        candidate_case_numbers=candidates,
        participants=json.loads(m.participants or "[]"),
        case_background=m.case_background or "",
        audio_filename=m.audio_filename,
        total_size_bytes=m.total_size_bytes or 0,
        duration_seconds=m.duration_seconds or 0,
        state=m.state,
        error_message=m.error_message or "",
        created_at=m.created_at.isoformat() if m.created_at else "",
        updated_at=m.updated_at.isoformat() if m.updated_at else "",
    )


def _orm_to_summary(m: MeetingORM) -> MeetingSummaryResponse:
    try:
        candidates = json.loads(m.candidate_case_numbers or "[]")
    except Exception:
        candidates = []
    return MeetingSummaryResponse(
        id=m.id,
        title=m.title,
        meeting_type=m.meeting_type,
        participants=json.loads(m.participants or "[]"),
        case_background=m.case_background or "",
        transcript=m.transcript or "",
        segments=json.loads(m.segments or "[]"),
        summary=json.loads(m.summary or "{}"),
        case_id=m.case_id,
        case_match_status=m.case_match_status,
        candidate_case_numbers=candidates,
        created_at=m.created_at.isoformat() if m.created_at else "",
    )


def _orm_to_case_meeting(m: MeetingORM) -> CaseMeetingItem:
    summary = json.loads(m.summary or "{}")
    todos = summary.get("todos", [])
    risks = summary.get("risk_alerts", [])
    try:
        candidates = json.loads(m.candidate_case_numbers or "[]")
    except Exception:
        candidates = []
    return CaseMeetingItem(
        id=m.id,
        title=m.title,
        meeting_type=m.meeting_type,
        participants=json.loads(m.participants or "[]"),
        case_background=m.case_background or "",
        duration_seconds=m.duration_seconds or 0,
        todo_count=len(todos) if isinstance(todos, list) else 0,
        risk_count=len(risks) if isinstance(risks, list) else 0,
        candidate_case_numbers=candidates,
        case_id=m.case_id,
        case_match_status=m.case_match_status,
        state=m.state,
        created_at=m.created_at.isoformat() if m.created_at else "",
    )


def _get_db() -> Session:
    db = SessionLocal()
    try:
        return db
    except Exception:
        db.close()
        raise


def _validate_audio_type(filename: str) -> str:
    ext = Path(filename or "audio.wav").suffix.lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}. Allowed: {sorted(ALLOWED_EXT)}")
    return ext


def _parse_participants(raw: str) -> list[str]:
    if not raw:
        return []
    try:
        data = json.loads(raw)
        if isinstance(data, list):
            return [str(x) for x in data]
    except (json.JSONDecodeError, TypeError):
        pass
    return [s.strip() for s in raw.split(",") if s.strip()]


def _normalize_case_number(text: str) -> str:
    if not text:
        return ""
    text = text.strip()
    text = text.replace("（", "(").replace("）", ")")
    text = text.replace("〔", "(").replace("〕", ")")
    return text


def _broadcast_progress(task_id: str, progress: TaskProgress) -> None:
    task_progress_store[task_id] = progress
    ws_list = task_ws_connections.get(task_id, [])
    dead: list[WebSocket] = []
    msg = json.dumps({
        "task_id": progress.task_id,
        "state": progress.state,
        "transcribed_seconds": progress.transcribed_seconds,
        "total_seconds": progress.total_seconds,
        "percent": round(progress.transcribed_seconds / progress.total_seconds * 100, 1) if progress.total_seconds > 0 else 0,
        "error": progress.error,
    })
    for ws in ws_list:
        try:
            asyncio.get_event_loop().create_task(ws.send_text(msg))
        except Exception:
            dead.append(ws)
    for d in dead:
        ws_list.remove(d)


async def send_wechat_notify(meeting: MeetingORM, success: bool, error_message: str = "") -> None:
    if not WECHAT_WEBHOOK_URL:
        return
    try:
        summary = json.loads(meeting.summary or "{}")
        todos = summary.get("todos", [])
        risks = summary.get("risk_alerts", [])
        todo_count = len(todos) if isinstance(todos, list) else 0
        risk_count = len(risks) if isinstance(risks, list) else 0

        case_id = meeting.case_id
        is_matched = meeting.case_match_status == "matched" and meeting.case_id
        candidate_cases = json.loads(meeting.candidate_case_numbers or "[]")
        if is_matched:
            case_timeline_url = f"{VIEW_BASE_URL}/case/{case_id}"
            case_line = f"> 案件编号：[{case_id}]({case_timeline_url})（案件时间线）"
        elif candidate_cases:
            case_info = "待关联（" + "、".join(candidate_cases) + "）"
            case_line = f"> 案件编号：{case_info}"
        else:
            case_line = "> 案件编号：待关联"

        view_url = f"{VIEW_BASE_URL}/meeting/{meeting.id}"
        status_text = "✅ 已生成" if success else "❌ 生成失败"
        content_lines = [
            f"**会议纪要{status_text}**",
            f"> 会议标题：{meeting.title}",
            case_line,
            f"> 查看纪要：[点击查看]({view_url})",
        ]

        if success:
            content_lines.append(f"> 待办事项：**{todo_count}** 条")
            content_lines.append(f"> 风险提示：**{risk_count}** 条")
            if meeting.duration_seconds:
                mins = int(meeting.duration_seconds // 60)
                secs = int(meeting.duration_seconds % 60)
                content_lines.append(f"> 会议时长：{mins}分{secs}秒")
        else:
            content_lines.append(f"> 失败原因：{error_message[:200]}")
            content_lines.append(f"> 会议ID：`{meeting.id}`（可用于后台重试/排查）")

        payload = {
            "msgtype": "markdown",
            "markdown": {"content": "\n".join(content_lines)},
        }
        async with httpx.AsyncClient(timeout=10.0) as client:
            await client.post(WECHAT_WEBHOOK_URL, json=payload)
    except Exception:
        pass


async def _process_task(meeting_id: str) -> None:
    db = _get_db()
    try:
        meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
        if not meeting:
            return

        progress = TaskProgress(task_id=meeting_id, state="transcribing", total_seconds=0.0)
        _broadcast_progress(meeting_id, progress)

        meeting.state = "transcribing"
        db.commit()

        audio_path = os.path.join(UPLOAD_DIR, meeting.audio_filename)
        if not os.path.exists(audio_path):
            meeting.state = "failed"
            meeting.error_message = "Audio file not found"
            db.commit()
            progress.state = "failed"
            progress.error = "Audio file not found"
            _broadcast_progress(meeting_id, progress)
            await send_wechat_notify(meeting, False, meeting.error_message)
            return

        existing_cases = []
        try:
            rows = db.execute(text("SELECT DISTINCT case_id FROM meetings WHERE case_id IS NOT NULL AND case_match_status = 'matched'")).fetchall()
            existing_cases = [{"case_id": r[0], "case_number": r[0]} for r in rows]
        except Exception:
            pass

        def on_progress(p: TaskProgress) -> None:
            _broadcast_progress(meeting_id, p)

        participants = json.loads(meeting.participants or "[]")
        result = await pipeline.run(
            audio_path=audio_path,
            meeting_type=MeetingType(meeting.meeting_type),
            participants=participants,
            case_background=meeting.case_background or None,
            on_progress=on_progress,
            existing_cases=existing_cases,
        )

        meeting.transcript = result["transcript"]
        meeting.segments = json.dumps(result["segments"], ensure_ascii=False)
        meeting.summary = json.dumps(result["summary"], ensure_ascii=False)
        meeting.case_id = result["summary"].get("case_id")
        meeting.case_match_status = result["summary"].get("case_match_status", "pending")
        candidates = result["summary"].get("candidate_case_numbers", [])
        meeting.candidate_case_numbers = json.dumps(candidates, ensure_ascii=False)
        meeting.state = "completed"
        meeting.updated_at = datetime.now()

        total_dur = sum(s["end"] - s["start"] for s in result["segments"]) if result["segments"] else 0
        meeting.duration_seconds = total_dur

        db.commit()

        progress.state = "completed"
        progress.transcribed_seconds = progress.total_seconds
        _broadcast_progress(meeting_id, progress)
        await send_wechat_notify(meeting, True)

    except Exception as exc:
        db = _get_db()
        meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
        error_msg = str(exc)
        if meeting:
            meeting.state = "failed"
            meeting.error_message = error_msg
            meeting.updated_at = datetime.now()
            db.commit()
        progress = TaskProgress(task_id=meeting_id, state="failed", error=error_msg)
        _broadcast_progress(meeting_id, progress)
        if meeting:
            await send_wechat_notify(meeting, False, error_msg)
    finally:
        db.close()


async def _queue_worker() -> None:
    while True:
        meeting_id = await task_queue.get()
        await _process_task(meeting_id)


@asynccontextmanager
async def lifespan(app: FastAPI):
    worker_task = asyncio.create_task(_queue_worker())
    yield
    worker_task.cancel()


app = FastAPI(title="会议录音转写与智能纪要生成API", version="1.4.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/api/upload", response_model=MeetingResponse)
async def upload_audio(
    file: UploadFile = File(...),
    title: str = Form(""),
    meeting_type: str = Form("client_consultation"),
    participants: str = Form("[]"),
    case_background: str = Form(""),
):
    ext = _validate_audio_type(file.filename)

    try:
        meeting_type_enum = MeetingType(meeting_type)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid meeting_type: {meeting_type}")

    participants_list = _parse_participants(participants)

    meeting_id = str(uuid.uuid4())
    filename = f"{meeting_id}{ext}"
    filepath = os.path.join(UPLOAD_DIR, filename)

    total_bytes = 0
    async with aiofiles.open(filepath, "wb") as f:
        while True:
            chunk = await file.read(1024 * 1024)
            if not chunk:
                break
            total_bytes += len(chunk)
            if total_bytes > MAX_FILE_SIZE_BYTES:
                try:
                    os.remove(filepath)
                except OSError:
                    pass
                raise HTTPException(status_code=400, detail=f"File too large. Max size: {MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB")
            await f.write(chunk)

    db = _get_db()
    meeting = MeetingORM(
        id=meeting_id,
        title=title or file.filename or "未命名会议",
        meeting_type=meeting_type_enum.value,
        participants=json.dumps(participants_list, ensure_ascii=False),
        case_background=case_background or "",
        audio_filename=filename,
        total_size_bytes=total_bytes,
        state="queued",
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )
    db.add(meeting)
    db.commit()
    db.refresh(meeting)

    progress = TaskProgress(task_id=meeting_id, state="queued", total_seconds=0.0)
    _broadcast_progress(meeting_id, progress)

    await task_queue.put(meeting_id)

    db.close()
    return _orm_to_response(meeting)


@app.get("/api/upload/chunks/{upload_id}")
async def get_upload_status(upload_id: str):
    chunk_dir = os.path.join(CHUNK_DIR, upload_id)
    manifest_path = os.path.join(chunk_dir, "manifest.json")
    if not os.path.exists(manifest_path):
        raise HTTPException(status_code=404, detail="Upload session not found")
    async with aiofiles.open(manifest_path, "r") as f:
        manifest = json.loads(await f.read())

    received_files: list[int] = []
    if os.path.isdir(chunk_dir):
        for name in os.listdir(chunk_dir):
            if name.startswith("chunk_"):
                try:
                    received_files.append(int(name.replace("chunk_", "")))
                except ValueError:
                    pass

    try:
        parts_list = json.loads(manifest.get("participants", "[]"))
    except (json.JSONDecodeError, TypeError):
        parts_list = []

    return {
        "upload_id": upload_id,
        "total_chunks": manifest.get("total_chunks", 0),
        "received_chunks": sorted(manifest.get("received", [])),
        "total_size_bytes": manifest.get("total_size_bytes", 0),
        "filename": manifest.get("filename", ""),
        "title": manifest.get("title", ""),
        "meeting_type": manifest.get("meeting_type", ""),
        "participants": parts_list,
        "case_background": manifest.get("case_background", ""),
    }


@app.post("/api/upload/chunk")
async def upload_chunk(
    file: UploadFile = File(...),
    upload_id: str = Form(""),
    chunk_index: int = Form(0),
    total_chunks: int = Form(1),
    total_size_bytes: int = Form(0),
    filename: str = Form("audio.wav"),
    title: str = Form(""),
    meeting_type: str = Form("client_consultation"),
    participants: str = Form("[]"),
    case_background: str = Form(""),
):
    if not upload_id:
        upload_id = str(uuid.uuid4())

    ext = _validate_audio_type(filename)

    if chunk_index < 0 or (total_chunks > 0 and chunk_index >= total_chunks):
        raise HTTPException(status_code=400, detail="Invalid chunk_index")

    if total_size_bytes and total_size_bytes > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=400, detail=f"Total file too large. Max size: {MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB")

    try:
        meeting_type_enum = MeetingType(meeting_type)
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Invalid meeting_type: {meeting_type}")

    participants_list = _parse_participants(participants)

    chunk_dir = os.path.join(CHUNK_DIR, upload_id)
    os.makedirs(chunk_dir, exist_ok=True)

    manifest_path = os.path.join(chunk_dir, "manifest.json")
    manifest: dict[str, Any] = {}
    if os.path.exists(manifest_path):
        async with aiofiles.open(manifest_path, "r") as f:
            manifest = json.loads(await f.read())
    else:
        manifest = {
            "upload_id": upload_id,
            "total_chunks": total_chunks,
            "received": [],
            "total_size_bytes": total_size_bytes,
            "filename": filename,
            "ext": ext,
            "title": title,
            "meeting_type": meeting_type_enum.value,
            "participants": json.dumps(participants_list, ensure_ascii=False),
            "case_background": case_background,
        }

    chunk_path = os.path.join(chunk_dir, f"chunk_{chunk_index:06d}")

    if chunk_index in manifest.get("received", []):
        if not os.path.exists(chunk_path):
            async with aiofiles.open(chunk_path, "wb") as f:
                content = await file.read()
                await f.write(content)
        async with aiofiles.open(manifest_path, "w") as f:
            await f.write(json.dumps(manifest, ensure_ascii=False))
        return {
            "status": "partial",
            "upload_id": upload_id,
            "received_chunks": len(manifest.get("received", [])),
            "total_chunks": manifest.get("total_chunks", total_chunks),
            "duplicate": True,
        }

    async with aiofiles.open(chunk_path, "wb") as f:
        content = await file.read()
        await f.write(content)

    manifest.setdefault("received", [])
    if chunk_index not in manifest["received"]:
        manifest["received"].append(chunk_index)
    manifest["received"] = sorted(manifest["received"])

    async with aiofiles.open(manifest_path, "w") as f:
        await f.write(json.dumps(manifest, ensure_ascii=False))

    if len(manifest["received"]) >= manifest.get("total_chunks", total_chunks):
        meeting_id = str(uuid.uuid4())
        final_filename = f"{meeting_id}{manifest.get('ext', ext)}"
        final_path = os.path.join(UPLOAD_DIR, final_filename)
        total_written = 0

        async with aiofiles.open(final_path, "wb") as outf:
            for i in range(manifest["total_chunks"]):
                cp = os.path.join(chunk_dir, f"chunk_{i:06d}")
                if not os.path.exists(cp):
                    try:
                        os.remove(final_path)
                    except OSError:
                        pass
                    raise HTTPException(status_code=400, detail=f"Missing chunk {i}")
                async with aiofiles.open(cp, "rb") as inf:
                    data = await inf.read()
                    total_written += len(data)
                    await outf.write(data)

        if total_written > MAX_FILE_SIZE_BYTES:
            try:
                os.remove(final_path)
            except OSError:
                pass
            raise HTTPException(status_code=400, detail=f"Total merged file too large. Max size: {MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB")

        shutil.rmtree(chunk_dir, ignore_errors=True)

        final_meeting_type = manifest.get("meeting_type", meeting_type)
        try:
            final_meeting_type_enum = MeetingType(final_meeting_type)
        except ValueError:
            final_meeting_type_enum = MeetingType.CLIENT_CONSULTATION

        try:
            final_participants = json.loads(manifest.get("participants", "[]"))
        except (json.JSONDecodeError, TypeError):
            final_participants = []

        db = _get_db()
        meeting = MeetingORM(
            id=meeting_id,
            title=manifest.get("title") or filename or "未命名会议",
            meeting_type=final_meeting_type_enum.value,
            participants=json.dumps(final_participants, ensure_ascii=False),
            case_background=manifest.get("case_background", "") or "",
            audio_filename=final_filename,
            total_size_bytes=total_written,
            state="queued",
            created_at=datetime.now(),
            updated_at=datetime.now(),
        )
        db.add(meeting)
        db.commit()
        db.refresh(meeting)

        progress = TaskProgress(task_id=meeting_id, state="queued")
        _broadcast_progress(meeting_id, progress)

        await task_queue.put(meeting_id)
        db.close()

        return _orm_to_response(meeting)

    return {
        "status": "partial",
        "upload_id": upload_id,
        "received_chunks": len(manifest["received"]),
        "total_chunks": manifest.get("total_chunks", total_chunks),
        "duplicate": False,
    }


@app.get("/api/meeting/{meeting_id}", response_model=MeetingResponse)
async def get_meeting(meeting_id: str):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    db.close()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    return _orm_to_response(meeting)


@app.get("/api/meeting/{meeting_id}/summary", response_model=MeetingSummaryResponse)
async def get_meeting_summary(meeting_id: str):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    db.close()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    if meeting.state != "completed":
        raise HTTPException(status_code=400, detail=f"Meeting is in state: {meeting.state}, not completed yet")
    return _orm_to_summary(meeting)


@app.get("/api/meetings", response_model=SearchResponse)
async def search_meetings(
    case_id: Optional[str] = None,
    date_from: Optional[str] = None,
    date_to: Optional[str] = None,
    participant: Optional[str] = None,
    keyword: Optional[str] = None,
    meeting_type: Optional[str] = None,
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
):
    db = _get_db()
    query = db.query(MeetingORM)

    if meeting_type:
        query = query.filter(MeetingORM.meeting_type == meeting_type)
    if date_from:
        query = query.filter(MeetingORM.created_at >= datetime.fromisoformat(date_from))
    if date_to:
        query = query.filter(MeetingORM.created_at <= datetime.fromisoformat(date_to))
    if participant:
        query = query.filter(MeetingORM.participants.contains(participant))
    if keyword:
        query = query.filter(
            (MeetingORM.transcript.contains(keyword))
            | (MeetingORM.summary.contains(keyword))
            | (MeetingORM.title.contains(keyword))
            | (MeetingORM.case_background.contains(keyword))
        )

    all_meetings = query.order_by(MeetingORM.created_at.desc()).all()

    if case_id:
        norm_query = _normalize_case_number(case_id)
        filtered: list[MeetingORM] = []
        for m in all_meetings:
            if m.case_match_status == "matched" and m.case_id:
                if _normalize_case_number(m.case_id) == norm_query:
                    filtered.append(m)
            else:
                try:
                    cands = json.loads(m.candidate_case_numbers or "[]")
                    for c in cands:
                        if _normalize_case_number(c) == norm_query:
                            filtered.append(m)
                            break
                except Exception:
                    pass
        all_meetings = filtered

    total = len(all_meetings)
    results = all_meetings[skip : skip + limit]
    db.close()

    return SearchResponse(
        results=[_orm_to_summary(m) for m in results],
        total=total,
    )


@app.get("/api/cases", response_model=CaseListResponse)
async def list_cases(
    status: Optional[str] = None,
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
):
    db = _get_db()
    query = db.query(MeetingORM).filter(MeetingORM.state == "completed")

    if status == "matched":
        query = query.filter(MeetingORM.case_match_status == "matched", MeetingORM.case_id.isnot(None))
    elif status == "unmatched":
        query = query.filter(MeetingORM.case_match_status == "unmatched")
    elif status == "pending":
        query = query.filter(MeetingORM.case_match_status == "pending")

    meetings = query.order_by(MeetingORM.created_at.desc()).all()

    cases_dict: dict[str, dict[str, Any]] = defaultdict(lambda: {
        "display_id": "",
        "meeting_count": 0,
        "completed_count": 0,
        "todo_count": 0,
        "risk_count": 0,
        "latest_meeting_time": None,
        "latest_meeting_title": "",
        "case_match_status": "pending",
    })

    for m in meetings:
        is_matched = m.case_match_status == "matched" and m.case_id

        case_ids_for_this_meeting: list[str] = []
        if is_matched:
            norm_id = _normalize_case_number(m.case_id)
            case_ids_for_this_meeting.append(norm_id)
            cases_dict[norm_id]["display_id"] = m.case_id
            cases_dict[norm_id]["case_match_status"] = "matched"
        else:
            if m.candidate_case_numbers:
                try:
                    cands = json.loads(m.candidate_case_numbers)
                    for c in cands:
                        if not c:
                            continue
                        norm_c = _normalize_case_number(c)
                        if norm_c not in case_ids_for_this_meeting:
                            case_ids_for_this_meeting.append(norm_c)
                            if not cases_dict[norm_c]["display_id"]:
                                cases_dict[norm_c]["display_id"] = c
                            if m.case_match_status == "unmatched":
                                cases_dict[norm_c]["case_match_status"] = "unmatched"
                except Exception:
                    pass

        if not case_ids_for_this_meeting:
            case_ids_for_this_meeting.append("__unclassified__")
            cases_dict["__unclassified__"]["display_id"] = ""

        for cid in case_ids_for_this_meeting:
            cases_dict[cid]["meeting_count"] += 1
            if m.state == "completed":
                cases_dict[cid]["completed_count"] += 1

            try:
                summary = json.loads(m.summary or "{}")
                todos = summary.get("todos", [])
                risks = summary.get("risk_alerts", [])
                cases_dict[cid]["todo_count"] += len(todos) if isinstance(todos, list) else 0
                cases_dict[cid]["risk_count"] += len(risks) if isinstance(risks, list) else 0
            except Exception:
                pass

            if cases_dict[cid]["latest_meeting_time"] is None or (m.created_at and m.created_at > cases_dict[cid]["latest_meeting_time"]):
                cases_dict[cid]["latest_meeting_time"] = m.created_at
                cases_dict[cid]["latest_meeting_title"] = m.title

    case_items: list[CaseSummaryItem] = []
    for norm_id, data in cases_dict.items():
        case_items.append(CaseSummaryItem(
            case_id=data["display_id"],
            case_match_status=data["case_match_status"],
            meeting_count=data["meeting_count"],
            completed_count=data["completed_count"],
            todo_count=data["todo_count"],
            risk_count=data["risk_count"],
            latest_meeting_time=data["latest_meeting_time"].isoformat() if data["latest_meeting_time"] else "",
            latest_meeting_title=data["latest_meeting_title"],
        ))

    case_items.sort(key=lambda x: x.latest_meeting_time or "", reverse=True)

    total = len(case_items)
    paged = case_items[skip : skip + limit]
    db.close()

    return CaseListResponse(cases=paged, total=total)


@app.get("/api/cases/{case_id}/meetings", response_model=CaseTimelineResponse)
async def get_case_meetings(
    case_id: str,
    meeting_type: Optional[str] = None,
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
):
    db = _get_db()
    query = db.query(MeetingORM)

    if meeting_type:
        query = query.filter(MeetingORM.meeting_type == meeting_type)

    norm_query = _normalize_case_number(case_id)
    meetings = query.order_by(MeetingORM.created_at.desc()).all()

    matched_meetings: list[MeetingORM] = []
    for m in meetings:
        if m.case_match_status == "matched" and m.case_id:
            if _normalize_case_number(m.case_id) == norm_query:
                matched_meetings.append(m)
        else:
            try:
                cands = json.loads(m.candidate_case_numbers or "[]")
                for c in cands:
                    if _normalize_case_number(c) == norm_query:
                        matched_meetings.append(m)
                        break
            except Exception:
                pass

    total = len(matched_meetings)
    paged = matched_meetings[skip : skip + limit]
    db.close()

    items = [_orm_to_case_meeting(m) for m in paged]
    return CaseTimelineResponse(case_id=case_id, meetings=items, total=total)


@app.get("/api/meeting/{meeting_id}/search")
async def search_in_meeting(meeting_id: str, keyword: str = Query(..., min_length=1)):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    db.close()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    segments = json.loads(meeting.segments or "[]")
    summary_data = json.loads(meeting.summary or "{}")
    kw = keyword.lower()

    transcript_hits: list[dict[str, Any]] = []
    for seg in segments:
        if kw in seg.get("text", "").lower():
            transcript_hits.append({
                "source": "transcript",
                "start": seg["start"],
                "end": seg["end"],
                "speaker": seg.get("speaker", ""),
                "text": seg["text"],
            })

    summary_fields = [
        ("topic", "会议主题"),
        ("raw_summary", "纪要全文"),
    ]
    list_fields = [
        ("todos", "待办事项", "item"),
        ("client_demands", "客户诉求", None),
        ("risk_alerts", "风险提示", None),
        ("key_points", "关键讨论点", "content"),
    ]

    summary_hits: list[dict[str, Any]] = []

    for key, label in summary_fields:
        val = summary_data.get(key, "")
        if isinstance(val, str) and kw in val.lower():
            summary_hits.append({
                "source": "summary",
                "field": key,
                "field_label": label,
                "text": val,
            })

    for key, label, item_key in list_fields:
        items = summary_data.get(key, [])
        if not isinstance(items, list):
            continue
        for idx, it in enumerate(items):
            if isinstance(it, dict):
                text = it.get(item_key, "") if item_key else json.dumps(it, ensure_ascii=False)
            else:
                text = str(it)
            if kw in text.lower():
                summary_hits.append({
                    "source": "summary",
                    "field": key,
                    "field_label": label,
                    "index": idx,
                    "text": text,
                })

    if isinstance(summary_data.get("raw_summary", ""), str):
        raw = summary_data["raw_summary"]
        lines = raw.split("\n")
        for idx, line in enumerate(lines):
            if kw in line.lower():
                already = any(
                    h.get("field") == "raw_summary" and h.get("text") == line.strip()
                    for h in summary_hits
                )
                if not already and line.strip():
                    summary_hits.append({
                        "source": "summary",
                        "field": "raw_summary",
                        "field_label": "纪要原文",
                        "line_index": idx,
                        "text": line.strip(),
                    })

    all_hits = transcript_hits + summary_hits

    return {
        "meeting_id": meeting_id,
        "keyword": keyword,
        "hits": all_hits,
        "total_hits": len(all_hits),
        "transcript_hits": len(transcript_hits),
        "summary_hits": len(summary_hits),
    }


@app.websocket("/ws/task/{task_id}")
async def websocket_task_progress(websocket: WebSocket, task_id: str):
    await websocket.accept()
    if task_id not in task_ws_connections:
        task_ws_connections[task_id] = []
    task_ws_connections[task_id].append(websocket)

    if task_id in task_progress_store:
        progress = task_progress_store[task_id]
        msg = json.dumps({
            "task_id": progress.task_id,
            "state": progress.state,
            "transcribed_seconds": progress.transcribed_seconds,
            "total_seconds": progress.total_seconds,
            "percent": round(progress.transcribed_seconds / progress.total_seconds * 100, 1) if progress.total_seconds > 0 else 0,
            "error": progress.error,
        })
        await websocket.send_text(msg)

    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        pass
    finally:
        if task_id in task_ws_connections:
            if websocket in task_ws_connections[task_id]:
                task_ws_connections[task_id].remove(websocket)
            if not task_ws_connections[task_id]:
                del task_ws_connections[task_id]


@app.get("/api/meeting/{meeting_id}/export/word")
async def export_word(meeting_id: str):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    db.close()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    if meeting.state != "completed":
        raise HTTPException(status_code=400, detail="Meeting not completed yet")

    from docx import Document

    doc = Document()
    doc.add_heading("会议纪要", level=0)

    info_table = doc.add_table(rows=6, cols=2)
    if meeting.case_match_status == "matched" and meeting.case_id:
        case_display = meeting.case_id
    else:
        case_display = "待关联"
        try:
            cands = json.loads(meeting.candidate_case_numbers or "[]")
            if cands:
                case_display = "待关联（识别到：" + "、".join(cands) + "）"
        except Exception:
            pass

    info_data = [
        ("会议主题", meeting.title),
        ("会议类型", meeting.meeting_type),
        ("参会人员", ", ".join(json.loads(meeting.participants or "[]"))),
        ("会议时间", meeting.created_at.isoformat() if meeting.created_at else ""),
        ("案件编号", case_display),
        ("案件背景", meeting.case_background or "无"),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value

    doc.add_heading("转写文本", level=1)
    doc.add_paragraph(meeting.transcript or "")

    summary_data = json.loads(meeting.summary or "{}")
    doc.add_heading("智能纪要", level=1)
    doc.add_paragraph(summary_data.get("raw_summary", ""))

    export_dir = os.path.join(UPLOAD_DIR, "exports")
    os.makedirs(export_dir, exist_ok=True)
    out_path = os.path.join(export_dir, f"{meeting_id}.docx")
    doc.save(out_path)

    return FileResponse(
        out_path,
        filename=f"{meeting.title or '会议纪要'}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.get("/api/meeting/{meeting_id}/export/pdf")
async def export_pdf(meeting_id: str):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    db.close()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")
    if meeting.state != "completed":
        raise HTTPException(status_code=400, detail="Meeting not completed yet")

    if meeting.case_match_status == "matched" and meeting.case_id:
        case_display = meeting.case_id
    else:
        case_display = "待关联"
        try:
            cands = json.loads(meeting.candidate_case_numbers or "[]")
            if cands:
                case_display = "待关联（识别到：" + "、".join(cands) + "）"
        except Exception:
            pass

    summary_data = json.loads(meeting.summary or "{}")
    html_content = f"""
    <html><head><meta charset="utf-8"><style>
    body {{ font-family: SimSun, serif; margin: 40px; }}
    h1 {{ text-align: center; }}
    table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
    td, th {{ border: 1px solid #000; padding: 8px; }}
    h2 {{ border-bottom: 2px solid #333; padding-bottom: 5px; }}
    pre {{ white-space: pre-wrap; }}
    </style></head><body>
    <h1>会议纪要</h1>
    <table>
    <tr><td><strong>会议主题</strong></td><td>{meeting.title}</td></tr>
    <tr><td><strong>会议类型</strong></td><td>{meeting.meeting_type}</td></tr>
    <tr><td><strong>参会人员</strong></td><td>{', '.join(json.loads(meeting.participants or '[]'))}</td></tr>
    <tr><td><strong>会议时间</strong></td><td>{meeting.created_at.isoformat() if meeting.created_at else ''}</td></tr>
    <tr><td><strong>案件编号</strong></td><td>{case_display}</td></tr>
    <tr><td><strong>案件背景</strong></td><td>{meeting.case_background or '无'}</td></tr>
    </table>
    <h2>转写文本</h2>
    <pre>{meeting.transcript or ''}</pre>
    <h2>智能纪要</h2>
    <pre>{summary_data.get('raw_summary', '')}</pre>
    </body></html>
    """

    from weasyprint import HTML
    export_dir = os.path.join(UPLOAD_DIR, "exports")
    os.makedirs(export_dir, exist_ok=True)
    out_path = os.path.join(export_dir, f"{meeting_id}.pdf")
    HTML(string=html_content).write_pdf(out_path)

    return FileResponse(out_path, filename=f"{meeting.title or '会议纪要'}.pdf", media_type="application/pdf")


@app.put("/api/meeting/{meeting_id}/case", response_model=MeetingResponse)
async def update_meeting_case(meeting_id: str, case_id: str = "", case_match_status: str = "matched"):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    if not meeting:
        db.close()
        raise HTTPException(status_code=404, detail="Meeting not found")
    meeting.case_id = case_id
    meeting.case_match_status = case_match_status
    meeting.updated_at = datetime.now()
    db.commit()
    resp = _orm_to_response(meeting)
    db.close()
    return resp


@app.put("/api/meeting/{meeting_id}/transcript")
async def update_meeting_transcript(meeting_id: str, transcript: str = "", summary_raw: str = ""):
    db = _get_db()
    meeting = db.query(MeetingORM).filter(MeetingORM.id == meeting_id).first()
    if not meeting:
        db.close()
        raise HTTPException(status_code=404, detail="Meeting not found")
    if transcript:
        meeting.transcript = transcript
    if summary_raw:
        current = json.loads(meeting.summary or "{}")
        current["raw_summary"] = summary_raw
        meeting.summary = json.dumps(current, ensure_ascii=False)
    meeting.updated_at = datetime.now()
    db.commit()
    db.close()
    return {"message": "Transcript updated", "meeting_id": meeting_id}


@app.get("/api/templates", response_model=list[TemplateRecord])
async def list_templates(meeting_type: Optional[str] = None):
    mt = MeetingType(meeting_type) if meeting_type else None
    return template_manager.list_all(mt)


@app.post("/api/templates", response_model=TemplateRecord)
async def create_template(req: TemplateCreateRequest):
    try:
        return template_manager.create(req)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))


@app.put("/api/templates/{template_id}", response_model=TemplateRecord)
async def update_template(template_id: str, req: TemplateUpdateRequest):
    result = template_manager.update(template_id, req)
    if result is None:
        raise HTTPException(status_code=404, detail="Template not found")
    return result


@app.delete("/api/templates/{template_id}")
async def delete_template(template_id: str):
    if not template_manager.delete(template_id):
        raise HTTPException(status_code=404, detail="Template not found")
    return {"message": "Template deleted"}


@app.get("/api/templates/{template_id}/versions/{version}")
async def get_template_version(template_id: str, version: int):
    record = template_manager.get_version(template_id, version)
    if record is None:
        raise HTTPException(status_code=404, detail="Version not found")
    return record


@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "timestamp": datetime.now().isoformat(),
        "version": "1.4.0",
        "wechat_webhook_configured": bool(WECHAT_WEBHOOK_URL),
        "max_file_size_mb": MAX_FILE_SIZE_BYTES // (1024 * 1024),
        "view_base_url": VIEW_BASE_URL,
    }


if __name__ == "__main__":
    uvicorn.run("server.main:app", host="0.0.0.0", port=8000, reload=True)
